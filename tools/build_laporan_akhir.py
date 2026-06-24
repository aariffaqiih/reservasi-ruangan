from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "laporan" / "Laporan Akhir PBO.docx"
CLASS_DIAGRAM = ROOT / "class-diagram.png"
UI_DIR = ROOT / "ui"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths):
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = Inches(width)
            set_cell_margins(row.cells[idx])
            row.cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_page_number(section):
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run()
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    run._r.append(fld)


def style_doc(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    add_page_number(section)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.333

    for name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 18, 10),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def add_centered(doc, text, size=12, bold=False, after=6):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(after)
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)
    return p


def add_cover(doc):
    add_centered(doc, "LAPORAN AKHIR TUGAS BESAR", 16, True, 4)
    add_centered(doc, "RESERVASI RUANGAN DI LUAR JAM KERJA", 15, True, 18)
    add_centered(doc, "Mata Kuliah:", 12, False, 0)
    add_centered(doc, "Pemrograman Berorientasi Objek", 12, True, 12)
    add_centered(doc, "Dosen Pengampu:", 12, False, 0)
    add_centered(doc, "Dany Candra Febrianto", 12, True, 20)
    add_centered(doc, "Nama Kelompok: Kicau Mania", 12, True, 12)
    add_centered(doc, "Disusun oleh:", 12, False, 8)
    for name in [
        "Amelia Sofiana Makharomi - 103112400233",
        "Ali Abdul Fattah 'Alim Kautsar - 103112400213",
        "Aarif Rahmaan Jalaluddin Faqiih - 103112430182",
        "Anggota Kelompok Kicau Mania",
    ]:
        add_centered(doc, name, 11, False, 2)
    doc.add_paragraph()
    doc.add_paragraph()
    add_centered(doc, "FAKULTAS INFORMATIKA", 12, True, 0)
    add_centered(doc, "TELKOM UNIVERSITY PURWOKERTO", 12, True, 0)
    add_centered(doc, "2026", 12, True, 0)
    doc.add_page_break()


def add_toc(doc):
    doc.add_heading("Daftar Isi", level=1)
    entries = [
        "BAB I Pendahuluan",
        "1.1 Latar Belakang",
        "1.2 Rumusan Masalah",
        "1.3 Batasan Masalah",
        "1.4 Tujuan",
        "1.5 Manfaat",
        "1.6 Penjelasan Singkat Aplikasi",
        "BAB II Pembahasan",
        "2.1 Analisis Arsitektur Sistem",
        "2.2 Pemodelan Entitas Sistem",
        "2.3 Implementasi Pilar PBO Pada Program",
        "2.4 Konstruksi UI dan Fungsionalitas",
        "2.5 Rincian Distribusi Tugas",
        "BAB III Penutup",
        "Daftar Pustaka",
    ]
    for e in entries:
        p = doc.add_paragraph(e)
        p.style = "List Bullet" if e[0].isdigit() else doc.styles["Normal"]
    doc.add_page_break()


def add_bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbered(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Number")


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        set_cell_shading(hdr[i], "F4F6F9")
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
    set_table_width(table, widths)
    doc.add_paragraph()
    return table


def add_image(doc, path, caption, width=6.2):
    if not path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run()
    r.add_picture(str(path), width=Inches(width))
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.runs[0].italic = True


def build():
    doc = Document()
    style_doc(doc)
    add_cover(doc)
    add_toc(doc)

    doc.add_heading("BAB I", level=1)
    doc.add_heading("PENDAHULUAN", level=1)
    doc.add_heading("1.1 Latar Belakang", level=2)
    for text in [
        "Penggunaan ruang kelas di lingkungan kampus tidak hanya terjadi pada jam perkuliahan reguler. Mahasiswa sering membutuhkan ruangan untuk rapat organisasi, diskusi kelompok, praktikum tambahan, seminar kecil, persiapan lomba, dan kegiatan akademik lain di luar jadwal kelas. Jika proses peminjaman masih dilakukan secara manual, petugas administrasi harus memeriksa ketersediaan ruangan, mencatat pemohon, meminta persetujuan, dan menginformasikan status secara terpisah. Pola kerja tersebut rawan menimbulkan keterlambatan, benturan jadwal, dan kehilangan riwayat penggunaan ruangan.",
        "Berdasarkan kebutuhan tersebut, proyek ini mengembangkan aplikasi web Sistem Reservasi Ruangan di Luar Jam Kerja. Sistem dirancang untuk mempertemukan tiga aktor utama, yaitu Mahasiswa sebagai pengaju reservasi, Admin sebagai pihak verifikator, dan Satpam sebagai petugas yang mengonfirmasi akses ruang saat kegiatan berlangsung. Melalui alur terpusat, mahasiswa dapat mengajukan reservasi, admin dapat memberikan keputusan, dan satpam dapat mencatat check-in serta check-out pemakaian ruangan.",
        "Pengembangan sistem juga menjadi media penerapan konsep Pemrograman Berorientasi Objek. Konsep inheritance diterapkan pada hierarki User, encapsulation diterapkan melalui pembatasan akses data dan penggunaan method, polymorphism diterapkan melalui interface Notifiable, sedangkan relasi antarobjek dibangun melalui entitas Reservation, Room, Approval, AccessRecord, dan Notification. Dengan demikian, aplikasi tidak hanya menyelesaikan masalah administrasi, tetapi juga menunjukkan implementasi prinsip PBO dalam proyek nyata berbasis Java dan Spring Boot.",
    ]:
        doc.add_paragraph(text)

    doc.add_heading("1.2 Rumusan Masalah", level=2)
    add_numbered(doc, [
        "Bagaimana merancang sistem reservasi ruang kelas berbasis web yang dapat digunakan mahasiswa untuk mengajukan peminjaman ruangan secara daring?",
        "Bagaimana menerapkan konsep Pemrograman Berorientasi Objek dalam sistem reservasi menggunakan Java, Spring Boot, dan struktur class yang saling berelasi?",
        "Bagaimana mengelola proses pengajuan, verifikasi, notifikasi, dan pemantauan reservasi dengan melibatkan peran Mahasiswa, Admin, dan Satpam?",
        "Bagaimana memastikan validasi ketersediaan ruangan agar tidak terjadi benturan jadwal pada ruangan dan waktu yang sama?",
        "Bagaimana menyimpan data reservasi, approval, akses, pengguna, ruangan, dan notifikasi ke dalam basis data secara terstruktur?",
    ])

    doc.add_heading("1.3 Batasan Masalah", level=2)
    add_numbered(doc, [
        "Sistem difokuskan pada reservasi ruang kelas kampus di luar jam perkuliahan atau untuk kegiatan akademik/kemahasiswaan.",
        "Aktor yang dicakup adalah Mahasiswa, Admin, dan Satpam. Peran lain seperti dosen, bagian keuangan, atau pengelola inventaris belum menjadi cakupan utama.",
        "Aplikasi dikembangkan sebagai aplikasi web lokal menggunakan Spring Boot, Thymeleaf, JPA, dan MySQL/MariaDB.",
        "Validasi utama difokuskan pada status aktif ruangan, tanggal reservasi, jam mulai, jam selesai, dan benturan jadwal dengan reservasi berstatus PENDING, APPROVED, atau ACTIVE.",
        "Sistem belum mencakup integrasi pembayaran, peminjaman barang tambahan, tanda tangan digital, atau integrasi kalender eksternal.",
    ])

    doc.add_heading("1.4 Tujuan", level=2)
    add_numbered(doc, [
        "Membangun aplikasi reservasi ruangan berbasis web yang memudahkan mahasiswa dalam mengajukan peminjaman ruang.",
        "Menerapkan prinsip PBO seperti abstract class, interface, inheritance, encapsulation, polymorphism, association, dan aggregation dalam desain sistem.",
        "Mengimplementasikan arsitektur MVC dengan pemisahan controller, service, repository, entity, DTO, template, dan database.",
        "Menyediakan alur persetujuan admin serta pencatatan akses oleh satpam agar proses pemakaian ruangan dapat dilacak.",
        "Menghasilkan sistem yang menyimpan riwayat reservasi, approval, access record, dan notifikasi secara konsisten di basis data.",
    ])

    doc.add_heading("1.5 Manfaat", level=2)
    doc.add_heading("1.5.1 Bagi Mahasiswa", level=3)
    add_bullets(doc, [
        "Mempermudah pengajuan reservasi tanpa harus datang langsung ke bagian administrasi.",
        "Memudahkan pemantauan status pengajuan, baik menunggu persetujuan, disetujui, ditolak, aktif, selesai, maupun dibatalkan.",
    ])
    doc.add_heading("1.5.2 Bagi Admin", level=3)
    add_bullets(doc, [
        "Membantu proses verifikasi permintaan reservasi berdasarkan data ruangan dan jadwal yang tercatat.",
        "Menyediakan riwayat keputusan berupa setuju, tolak, atau revisi beserta catatan admin.",
    ])
    doc.add_heading("1.5.3 Bagi Satpam", level=3)
    add_bullets(doc, [
        "Mempermudah pencatatan check-in dan check-out penggunaan ruangan.",
        "Membantu dokumentasi kendala atau pelanggaran yang terjadi selama kegiatan berlangsung.",
    ])
    doc.add_heading("1.5.4 Bagi Pengembang Akademik", level=3)
    add_bullets(doc, [
        "Menjadi bukti penerapan konsep PBO dalam studi kasus yang dekat dengan kebutuhan kampus.",
        "Melatih pembuatan aplikasi berlapis menggunakan Java, Spring Boot, ORM, template web, dan database relasional.",
    ])

    doc.add_heading("1.6 Penjelasan Singkat Aplikasi", level=2)
    doc.add_paragraph("Sistem Reservasi Ruangan di Luar Jam Kerja adalah aplikasi web yang berfungsi sebagai pusat pengelolaan peminjaman ruangan kampus. Pengguna mengakses aplikasi melalui halaman web, lalu sistem menampilkan fitur sesuai peran. Mahasiswa mengisi data reservasi berupa ruangan, tanggal, jam mulai, jam selesai, dan tujuan kegiatan. Sistem memvalidasi ketersediaan ruangan sebelum menyimpan pengajuan.")
    doc.add_paragraph("Admin melihat daftar pengajuan yang menunggu keputusan, memeriksa kelayakan jadwal, lalu memberikan keputusan setuju, tolak, atau revisi. Keputusan tersebut disimpan pada entitas Approval dan status reservasi diperbarui. Sistem juga membuat notifikasi agar mahasiswa memperoleh informasi perkembangan reservasi.")
    doc.add_paragraph("Pada hari pelaksanaan, Satpam dapat mengonfirmasi check-in ketika peminjam mulai menggunakan ruangan dan check-out ketika kegiatan selesai. Jika terdapat kendala, petugas dapat menambahkan catatan pada AccessRecord. Seluruh data tersimpan di database reservasi_ruang sehingga proses dapat ditelusuri kembali.")

    doc.add_page_break()
    doc.add_heading("BAB II", level=1)
    doc.add_heading("PEMBAHASAN", level=1)
    doc.add_heading("2.1 Analisis Arsitektur Sistem", level=2)
    doc.add_paragraph("Aplikasi dibangun menggunakan arsitektur Model-View-Controller. Lapisan View menggunakan Thymeleaf untuk halaman dashboard dan form, lapisan Controller menerima request web/REST, lapisan Service memuat logika bisnis, lapisan Repository menangani akses data melalui Spring Data JPA, sedangkan lapisan Entity merepresentasikan tabel database.")
    add_table(doc, ["Lapisan", "Komponen", "Peran"], [
        ("View", "templates dan static/css", "Menampilkan halaman dashboard mahasiswa, admin, satpam, daftar reservasi, form reservasi, profil, ruangan, approval, dan notifikasi."),
        ("Controller", "HomeController, AdminController, MahasiswaController, SatpamController, REST Controller", "Mengatur routing halaman dan endpoint API untuk operasi pengguna, ruangan, reservasi, approval, akses, dan notifikasi."),
        ("Service", "ReservationService, ApprovalService, AccessControlService, NotificationService, RoomService, UserService", "Menjalankan validasi, perubahan status, pengiriman notifikasi, dan koordinasi antarobjek."),
        ("Repository", "JpaRepository per entitas", "Mengambil dan menyimpan data ke database MySQL/MariaDB."),
        ("Model", "User, Mahasiswa, Admin, Satpam, Room, Reservation, Approval, AccessRecord, Notification", "Merepresentasikan domain sistem dan relasi objek sesuai prinsip PBO."),
    ], [1.1, 2.2, 3.2])
    doc.add_paragraph("Konfigurasi aplikasi menggunakan server port 8080, database reservasi_ruang, Hibernate ddl-auto update, dan dialect MySQL. Dengan struktur ini, sistem dapat dikembangkan secara modular karena perubahan pada tampilan, logika bisnis, dan akses data tidak tercampur dalam satu bagian kode.")

    doc.add_heading("2.2 Pemodelan Entitas Sistem", level=2)
    doc.add_heading("2.2.1 Identifikasi Entitas", level=3)
    add_table(doc, ["Class", "Peran", "Atribut Kunci", "Method Utama"], [
        ("User", "Abstract class induk pengguna.", "id, nama, email, noHp, passwordHash", "login(), logout(), ubahProfil()"),
        ("Mahasiswa", "Pengaju reservasi.", "nim, prodi, angkatan, reservations", "ajukanReservasi(), batalkanReservasi(), lihatStatusReservasi(), receiveNotification()"),
        ("Admin", "Verifikator pengajuan.", "unitKerja, approvals", "setujuiReservasi(), tolakReservasi(), mintaRevisiData(), receiveNotification()"),
        ("Satpam", "Petugas akses ruangan.", "shift, posJaga, accessRecords", "konfirmasiCheckIn(), konfirmasiCheckOut(), catatKendala(), receiveNotification()"),
        ("Room", "Data ruangan kampus.", "roomId, namaRuang, gedung, kapasitas, statusAktif", "aktifkan(), nonaktifkan(), ubahStatusAktif(), getInfoRuang()"),
        ("Reservation", "Data inti peminjaman.", "mahasiswa, room, tanggal, jamMulai, jamSelesai, tujuan, status", "ajukan(), ubahStatus(), batalkan(), validasiWaktu(), isCanBeCancelled()"),
        ("Approval", "Keputusan admin.", "reservation, admin, keputusan, catatan, reviewedAt", "setujui(), tolak(), mintaRevisi()"),
        ("AccessRecord", "Catatan penggunaan ruangan.", "reservation, satpam, checkInTime, checkOutTime, catatanPelanggaran", "checkIn(), checkOut(), laporkanKendala()"),
        ("Notification", "Pesan status untuk pengguna.", "penerima, reservation, pesan, statusBaca, createdAt", "kirim(), tandaiDibaca()"),
    ], [1.1, 1.7, 2.1, 1.6])

    doc.add_heading("2.2.2 Analisis Relasi Antar Objek", level=3)
    add_bullets(doc, [
        "Inheritance: Mahasiswa, Admin, dan Satpam merupakan turunan dari User. Strategi JPA yang digunakan adalah JOINED sehingga data umum berada pada tabel users dan data spesifik berada pada tabel turunan.",
        "Interface realization: Mahasiswa, Admin, dan Satpam mengimplementasikan Notifiable sehingga semua peran dapat menerima notifikasi dengan method receiveNotification().",
        "Association: Reservation terhubung dengan Mahasiswa dan Room. Approval terhubung dengan Reservation dan Admin. AccessRecord terhubung dengan Reservation dan Satpam.",
        "Aggregation: Service menyimpan kumpulan objek seperti reservations, rooms, approvals, accessRecords, dan notifications sebagai representasi pengelolaan domain.",
        "Enum: ReservationStatus menyimpan status DRAFT, PENDING, APPROVED, REJECTED, ACTIVE, COMPLETED, dan CANCELLED. ApprovalDecision menyimpan SETUJUI, TOLAK, dan REVISI.",
    ])
    add_image(doc, CLASS_DIAGRAM, "Gambar 2.1 Class diagram sistem reservasi ruangan.", 6.1)

    doc.add_heading("2.3 Implementasi Pilar PBO Pada Program", level=2)
    doc.add_heading("A. Enkapsulasi", level=3)
    doc.add_paragraph("Enkapsulasi terlihat dari pemodelan atribut pada class entity dan pengubahan data melalui method yang bermakna. Contohnya, Reservation tidak langsung diubah secara bebas pada alur bisnis utama, tetapi memiliki method ajukan(), ubahStatus(), batalkan(), validasiWaktu(), dan isCanBeCancelled(). User juga memiliki login(), logout(), dan ubahProfil() untuk mengatur data pengguna.")
    doc.add_heading("B. Abstraksi", level=3)
    doc.add_paragraph("Abstraksi diterapkan melalui class User sebagai representasi umum pengguna sistem. Detail spesifik seperti nim mahasiswa, unit kerja admin, serta shift dan pos jaga satpam ditempatkan pada subclass masing-masing. Lapisan service juga mengabstraksi logika bisnis sehingga controller tidak perlu mengetahui detail validasi database.")
    doc.add_heading("C. Polimorfisme", level=3)
    doc.add_paragraph("Polimorfisme diterapkan melalui interface Notifiable. Mahasiswa, Admin, dan Satpam memiliki bentuk implementasi receiveNotification() sesuai konteks objek masing-masing. Selain itu, penggunaan User sebagai tipe induk memungkinkan sistem mengambil data pengguna dari repository lalu mengecek tipe sebenarnya dengan instanceof sebelum menjalankan alur khusus Mahasiswa atau Admin.")
    doc.add_heading("D. Inheritance", level=3)
    doc.add_paragraph("Inheritance digunakan pada class Mahasiswa, Admin, dan Satpam yang mewarisi atribut dan perilaku dasar dari User. Pendekatan ini mengurangi duplikasi atribut seperti id, nama, email, noHp, dan passwordHash, sekaligus membuat model pengguna lebih konsisten.")

    doc.add_heading("2.4 Konstruksi UI dan Fungsionalitas", level=2)
    add_table(doc, ["Antarmuka", "Elemen Utama", "Fungsi", "Class Terkait"], [
        ("Halaman Awal/Login", "Navigasi dan akses peran", "Menghubungkan pengguna ke fitur sesuai peran.", "User, Mahasiswa, Admin, Satpam"),
        ("Dashboard Mahasiswa", "Profil, form reservasi, daftar status", "Mengajukan reservasi dan memantau status.", "Mahasiswa, Reservation, Room, Notification"),
        ("Dashboard Admin", "Daftar pending, tombol keputusan, catatan", "Memverifikasi dan memutuskan pengajuan.", "Admin, Approval, Reservation"),
        ("Dashboard Satpam", "Data reservasi aktif, check-in, check-out, catatan", "Mencatat akses penggunaan ruangan.", "Satpam, AccessRecord, Reservation"),
        ("Manajemen Ruangan", "Daftar ruangan, gedung, kapasitas, status aktif", "Melihat dan mengelola ketersediaan ruangan.", "Room, RoomService"),
        ("Notifikasi", "Daftar pesan dan status baca", "Memberi informasi perubahan status kepada pengguna.", "Notification, NotificationService"),
    ], [1.3, 1.8, 2.1, 1.3])
    ui_images = sorted(UI_DIR.glob("*.png"))[:4]
    for i, img in enumerate(ui_images, 1):
        add_image(doc, img, f"Gambar 2.{i+1} Tampilan antarmuka aplikasi.", 5.7)

    doc.add_heading("2.5 Rincian Distribusi Tugas", level=2)
    add_table(doc, ["Nama/Modul", "Tanggung Jawab", "Hasil Implementasi"], [
        ("Amelia Sofiana Makharomi", "Modul User dan Mahasiswa", "Abstract class User, subclass Mahasiswa, profil pengguna, dan alur pengajuan mahasiswa."),
        ("Ali Abdul Fattah 'Alim Kautsar", "ReservationService", "Validasi ketersediaan, pembuatan reservasi, pembatalan, riwayat, dan relasi reservasi."),
        ("Aarif Rahmaan Jalaluddin Faqiih", "Reservation dan ApprovalService", "Entity Reservation, status reservasi, verifikasi, keputusan approval, dan pengiriman notifikasi."),
        ("Kelompok Kicau Mania", "Integrasi aplikasi", "Controller, repository, template, CSS, database, class diagram, dan pengujian alur aplikasi."),
    ], [1.9, 2.0, 2.6])

    doc.add_page_break()
    doc.add_heading("BAB III", level=1)
    doc.add_heading("PENUTUP", level=1)
    doc.add_heading("3.1 Kesimpulan", level=2)
    doc.add_paragraph("Pengembangan Sistem Reservasi Ruangan di Luar Jam Kerja berhasil memodelkan kebutuhan peminjaman ruang kampus ke dalam aplikasi web berbasis Java Spring Boot. Sistem menyediakan alur utama mulai dari pengajuan oleh mahasiswa, verifikasi oleh admin, pengiriman notifikasi, hingga pencatatan akses oleh satpam. Data disimpan pada database relasional sehingga riwayat penggunaan ruangan dapat ditelusuri kembali.")
    doc.add_paragraph("Dari sisi Pemrograman Berorientasi Objek, proyek ini menerapkan inheritance melalui hierarki User, abstraction melalui pemisahan class induk dan service, encapsulation melalui method domain, polymorphism melalui interface Notifiable, serta association antarentitas Reservation, Room, Approval, AccessRecord, dan Notification. Implementasi tersebut menunjukkan bahwa konsep PBO dapat membantu membangun sistem yang lebih terstruktur, mudah dipahami, dan mudah dikembangkan.")
    doc.add_heading("3.2 Saran", level=2)
    add_numbered(doc, [
        "Menambahkan autentikasi dan otorisasi yang lebih lengkap agar setiap peran hanya dapat mengakses halaman sesuai haknya.",
        "Menambahkan fitur unggah surat izin atau dokumen pendukung reservasi.",
        "Mengintegrasikan kalender kegiatan agar jadwal ruangan dapat dilihat dalam bentuk visual bulanan atau mingguan.",
        "Menambahkan fitur laporan penggunaan ruangan berdasarkan gedung, tanggal, status, dan frekuensi peminjaman.",
        "Menyediakan pengujian otomatis untuk validasi benturan jadwal, perubahan status, dan pengiriman notifikasi.",
    ])

    doc.add_heading("Daftar Pustaka", level=1)
    for ref in [
        "Efan, Krismadinata, Jama, J., & Mulya, R. (2023). Object-Oriented Programming Course: A Systematic Literature Review on Problems and Solutions. International Journal of Information and Education Technology, 13(2), 1808-1815.",
        "Oracle. (2026). The Java Tutorials: Object-Oriented Programming Concepts. Oracle Java Documentation.",
        "Spring. (2026). Spring Boot Reference Documentation. VMware Tanzu Spring Documentation.",
        "Spring. (2026). Spring Data JPA Reference Documentation. VMware Tanzu Spring Documentation.",
        "Thymeleaf. (2026). Thymeleaf Documentation.",
        "MySQL. (2026). MySQL 8.0 Reference Manual.",
    ]:
        doc.add_paragraph(ref)

    for paragraph in doc.paragraphs:
        if paragraph.style.name == "Normal":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    OUT.parent.mkdir(exist_ok=True)
    doc.save(OUT)


if __name__ == "__main__":
    build()
